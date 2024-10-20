from docx import Document
from PyPDF2 import PdfReader

from datetime import date
import requests

import settings
from core import bot, BOT_TOKEN


def doc_parser(message) -> str:
    file_info = bot.get_file(message.document.file_id)
    download_url = f'https://api.telegram.org/file/bot{BOT_TOKEN}/{file_info.file_path}'

    r = requests.get(download_url, allow_redirects=True)

    if message.document.file_name.endswith('.pdf'):
        open('document.pdf', 'wb').write(r.content)

        pdf = PdfReader('document.pdf')
        pages = len(pdf.pages)
        text = ''
        for page_num in range(pages):
            page = pdf.pages[page_num]
            text += page.extract_text()
    else:
        open('document.doc', 'wb').write(r.content)
        doc = Document('document.doc')

        text = ''

        for para in doc.paragraphs:
            text += para.text + '\n'
    return text


def create_word_document(user_key: str, text: str) -> str:
    """
    Создает документ Word, в котором содержится предоставленный текст.
    Использует user_key для названия файла. Возвращает путь к созданному файлу.

    Args:
        user_key (str): ключ пользователя, используемый для создания названия файла.
        text (str): текст, который нужно добавить в документ.

    Returns:
        str: путь к созданному файлу .docx
    """

    name = settings.CURRENT_MODE.get_response(
        user_key,
        "Создай краткое название текста буквально в 3 словах и отправь мне"
    )

    doc = Document()
    doc.add_paragraph(text)
    file_path = f'{name}_Стенография документа_{date.today()}.docx'
    doc.save(file_path)
    return file_path
